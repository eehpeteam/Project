"""Meeting notes formatting module"""

import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.config import Config

logger = logging.getLogger(__name__)


class MeetingNoteFormatter:
    """Format and structure meeting notes"""
    
    def __init__(self, meeting_title="Team Meeting", participants=None):
        """
        Initialize note formatter
        
        Args:
            meeting_title (str): Title of the meeting
            participants (list): List of participant emails
        """
        self.meeting_title = meeting_title
        self.participants = participants or []
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Set up the document with header and formatting"""
        # Add title
        title = self.doc.add_heading(self.meeting_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        metadata_para = self.doc.add_paragraph()
        metadata_para.add_run('Meeting Date: ').bold = True
        metadata_para.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
        
        if self.participants:
            participants_para = self.doc.add_paragraph()
            participants_para.add_run('Participants: ').bold = True
            participants_para.add_run(", ".join(self.participants))
        
        self.doc.add_paragraph()  # Add spacing
    
    def add_section(self, section_title):
        """Add a section heading"""
        self.doc.add_heading(section_title, level=1)
    
    def add_text(self, text):
        """Add plain text paragraph"""
        self.doc.add_paragraph(text)
    
    def add_bullet_point(self, text, level=0):
        """Add a bullet point"""
        self.doc.add_paragraph(text, style=f'List Bullet' if level == 0 else f'List Bullet {level + 1}')
    
    def add_numbered_point(self, text, level=0):
        """Add a numbered point"""
        self.doc.add_paragraph(text, style=f'List Number' if level == 0 else f'List Number {level + 1}')
    
    def add_table(self, rows, cols, data=None):
        """
        Add a table to the document
        
        Args:
            rows (int): Number of rows
            cols (int): Number of columns
            data (list): List of lists containing table data
        """
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if data:
            for row_idx, row_data in enumerate(data):
                if row_idx < rows:
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < cols:
                            table.rows[row_idx].cells[col_idx].text = str(cell_text)
        
        return table
    
    def add_key_points_from_transcription(self, transcription_text):
        """
        Parse transcription and extract key points
        
        Args:
            transcription_text (str): Raw transcription text
        """
        self.add_section("Meeting Summary")
        
        # Split text into sentences and add as bullet points
        sentences = [s.strip() for s in transcription_text.split('.') if s.strip()]
        
        for sentence in sentences[:10]:  # Limit to first 10 sentences as key points
            if len(sentence) > 20:  # Filter out very short fragments
                self.add_bullet_point(sentence + ".")
        
        if len(sentences) > 10:
            self.add_paragraph(f"\n... ({len(sentences) - 10} more points in full transcription)")
    
    def add_full_transcription(self, transcription_text):
        """Add full transcription to document"""
        self.add_section("Full Transcription")
        self.add_text(transcription_text)
    
    def add_action_items(self, items=None):
        """
        Add action items section
        
        Args:
            items (list): List of action items
        """
        self.add_section("Action Items")
        
        if items:
            for idx, item in enumerate(items, 1):
                self.add_numbered_point(item)
        else:
            self.add_text("No specific action items identified from this meeting.")
    
    def add_next_meeting(self, date_time=None, topic=None):
        """
        Add next meeting information
        
        Args:
            date_time (str): Next meeting date and time
            topic (str): Topic for next meeting
        """
        self.add_section("Next Meeting")
        
        if date_time:
            para = self.doc.add_paragraph()
            para.add_run("Date & Time: ").bold = True
            para.add_run(date_time)
        
        if topic:
            para = self.doc.add_paragraph()
            para.add_run("Topic: ").bold = True
            para.add_run(topic)
    
    def save(self, filename=None):
        """
        Save the document
        
        Args:
            filename (str): Output filename. If None, generates timestamp-based name
            
        Returns:
            str: Path to saved document or None if save failed
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_notes_{timestamp}.docx"
        
        try:
            filepath = Path(Config.OUTPUT_DIRECTORY) / filename
            self.doc.save(str(filepath))
            logger.info(f"Meeting notes saved to {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return None
    
    def add_paragraph(self, text):
        """Add a paragraph with text"""
        self.doc.add_paragraph(text)
